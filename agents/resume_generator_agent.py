import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .content_selector_agent import TailoredResumeContent
from typing import Dict

def add_bottom_border_to_paragraph(paragraph):
    """
    This function directly manipulates the paragraph's XML to add a
    bottom border, mimicking the user's working document.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    
    pBdr.append(bottom)
    pPr.append(pBdr)

# --- Default Generator (Final Professional Version with all fixes) ---
def generate_resume_docx(
    tailored_content: TailoredResumeContent,
    contact_info: Dict,
    output_folder: str = "output"
) -> str:
    print("Generating .docx resume using XML-derived styling...")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # --- HEADER ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    run_name = p_name.add_run(contact_info.get('name', ''))
    run_name.font.size = Pt(14)
    run_name.font.bold = True
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(8)
    p_contact.add_run(contact_info.get('address', ''))
    p_contact.add_run().add_break()
    p_contact.add_run(f"{contact_info.get('email', '')} | {contact_info.get('phone', '')} | {contact_info.get('linkedin', '')}")
    
    def add_section_heading(text, first_section=False):
        p = doc.add_paragraph()
        if not first_section:
            p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        add_bottom_border_to_paragraph(p)

    # --- BUILD DOCUMENT SECTIONS ---
    add_section_heading('PROFESSIONAL SUMMARY', first_section=True)
    doc.add_paragraph(tailored_content.professional_summary)

    add_section_heading('EXPERIENCE')
    for job in tailored_content.selected_experience:
        p_job_title = doc.add_paragraph()
        p_job_title.paragraph_format.space_before = Pt(6)
        run_job = p_job_title.add_run(f"{job.role} – {job.company}")
        run_job.bold = True
        p_job_title.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run_date = p_job_title.add_run('\t' + job.dates)
        run_date.bold = True
        for point in job.rewritten_bullet_points:
            p_bullet = doc.add_paragraph(point, style='List Bullet')
            p_bullet.paragraph_format.space_before = Pt(6)

    add_section_heading('SKILLS')
    for category, skills_list in tailored_content.relevant_skills.items():
        p_skill = doc.add_paragraph(style='List Bullet')
        p_skill.add_run(f"{category.replace('_', ' ').title()}: ").bold = True
        p_skill.add_run(', '.join(skills_list))

    if tailored_content.selected_projects:
        add_section_heading('PROJECTS')
        for project in tailored_content.selected_projects:
            doc.add_paragraph(f"{project.name} - {project.rewritten_description}", style='List Bullet')
    
    if tailored_content.education:
        add_section_heading('EDUCATION')
        for edu in tailored_content.education:
            # Institution line
            p_inst = doc.add_paragraph()
            p_inst.paragraph_format.space_before = Pt(6)
            p_inst.paragraph_format.space_after = Pt(0) # <-- FIX: Tighten space
            p_inst.add_run(edu['institution']).bold = True
            p_inst.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            run_edu_date = p_inst.add_run('\t' + edu['dates'])
            run_edu_date.bold = True
            
            # Degree line
            p_degree = doc.add_paragraph()
            p_degree.paragraph_format.space_after = Pt(0) # <-- FIX: Tighten space
            p_degree.add_run(edu['degree']).italic = True
            if edu.get('gpa'): p_degree.add_run(f" (GPA: {edu.get('gpa')})")
            
            # Courses line (optional)
            if edu.get('relevant_courses'):
                p_courses = doc.add_paragraph()
                p_courses.paragraph_format.space_after = Pt(8) # <-- FIX: Space after entire entry
                p_courses.add_run("Relevant Courses: ").bold = True
                p_courses.add_run(', '.join(edu['relevant_courses']))
            else:
                # If no courses, add space after the degree line instead
                p_degree.paragraph_format.space_after = Pt(8)

    if tailored_content.accomplishments:
        add_section_heading('ACCOMPLISHMENTS')
        for acc in tailored_content.accomplishments:
            doc.add_paragraph(acc, style='List Bullet')

    company_name = tailored_content.selected_experience[0].company.replace(' ', '_').replace('/', '_')
    output_filename = f"Resume_For_{company_name}.docx"
    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)
    print(f"Successfully generated resume: {output_path}")
    return output_path

# --- Template-based Generator (Unchanged) ---
def find_and_replace(paragraph, placeholder, replacement_text):
    if placeholder in paragraph.text:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(replacement_text))
def replace_list_placeholder(doc, placeholder, replacement_list):
    for p in doc.paragraphs:
        if p.text.strip() == placeholder:
            style = p.style
            for item in reversed(replacement_list):
                new_p = doc.add_paragraph(str(item), style=style)
                p._p.addprevious(new_p._p)
            p._element.getparent().remove(p._element)
            return True
    return False
def generate_resume_from_template_docx(
    tailored_content: TailoredResumeContent,
    contact_info: Dict,
    template_path: str,
    output_folder: str = "output"
) -> str:
    print(f"Generating .docx resume from template: {template_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found at: {template_path}")
    doc = Document(template_path)
    contact_replacements = {'{{email}}': contact_info.get('email', ''),'{{phone}}': contact_info.get('phone', ''),'{{linkedin}}': contact_info.get('linkedin', ''),'{{github}}': contact_info.get('github', '')}
    text_replacements = {'{{professional_summary}}': tailored_content.professional_summary,'{{skills}}': ' • '.join([f"{cat.replace('_', ' ').title()}: {', '.join(skills)}" for cat, skills in tailored_content.relevant_skills.items()]),'{{education}}': ' | '.join([f"{edu['degree']} from {edu['institution']} ({edu['dates']})" for edu in tailored_content.education])}
    text_replacements.update(contact_replacements)
    for p in doc.paragraphs:
        for key, value in text_replacements.items():
            find_and_replace(p, key, value)
    replace_list_placeholder(doc, '{{accomplishments}}', tailored_content.accomplishments)
    for i, job in enumerate(tailored_content.selected_experience):
        replace_list_placeholder(doc, f'{{{{experience_{i}_bullets}}}}', job.rewritten_bullet_points)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    company_name = tailored_content.selected_experience[0].company.replace(' ', '_').replace('/', '_')
    output_filename = f"Resume_For_{company_name}_From_Template.docx"
    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)
    print(f"Successfully generated resume from template: {output_path}")
    return output_path